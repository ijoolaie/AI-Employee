"""Tenant-scoped document indexing and vector retrieval for RAG v0.2.23."""
from __future__ import annotations
import hashlib
import math
import re
import uuid
from pathlib import Path
import httpx
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.config import get_settings
from app.core.exceptions import NotFoundError, ValidationAppError
from app.core.logging import request_id_var
from app.rag.text import chunk_text

settings = get_settings()


def extract_text(file_obj) -> str:
    from app.services import storage
    backend = storage.get_storage_backend()
    with backend.open(file_obj.storage_key) as stream:
        raw = stream.read()
    content_type = (file_obj.content_type or "").lower()
    suffix = Path(file_obj.filename).suffix.lower()
    if content_type.startswith("text/") or suffix in {".txt", ".md", ".csv", ".json", ".xml", ".html"}:
        return raw.decode("utf-8", errors="replace")
    if suffix == ".pdf" or content_type == "application/pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ValidationAppError("PDF indexing requires the optional pypdf dependency") from exc
        import io
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx" or content_type.endswith("wordprocessingml.document"):
        try:
            from docx import Document
        except ImportError as exc:
            raise ValidationAppError("DOCX indexing requires the optional python-docx dependency") from exc
        import io
        doc = Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs)
    raise ValidationAppError(f"Unsupported knowledge file type: {file_obj.filename}")


def _deterministic_embedding(text: str, dimensions: int = 64) -> list[float]:
    """Stable lightweight embedding for Compose certification only."""
    values = [0.0] * dimensions
    tokens = re.findall(r"[a-z0-9]+", text.lower())
    for token in tokens:
        digest = hashlib.sha256(token.encode()).digest()
        index = int.from_bytes(digest[:2], "big") % dimensions
        sign = 1.0 if digest[2] & 1 else -1.0
        values[index] += sign
    norm = math.sqrt(sum(value * value for value in values))
    return [value / norm for value in values] if norm else [0.0] * dimensions


async def embed_texts(texts: list[str]) -> list[list[float]]:
    if not texts:
        return []
    if settings.e2e_deterministic_embeddings:
        return [_deterministic_embedding(text) for text in texts]
    url = f"{settings.lm_studio_base_url.rstrip('/')}/embeddings"
    headers = {"content-type": "application/json"}
    if settings.lm_studio_api_key:
        headers["authorization"] = f"Bearer {settings.lm_studio_api_key}"
    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(url, headers=headers, json={"model": settings.ai_embedding_model, "input": texts})
        response.raise_for_status()
        data = response.json()
    items = data.get("data") or []
    items = sorted(items, key=lambda item: int(item.get("index", 0)))
    embeddings = [item.get("embedding") for item in items]
    if len(embeddings) != len(texts) or any(not isinstance(v, list) for v in embeddings):
        raise RuntimeError("Embedding provider returned an invalid response")
    return embeddings


async def index_file(db: AsyncSession, *, tenant_id: uuid.UUID, file_id: uuid.UUID, actor_id: uuid.UUID):
    from app.models.file import FileObject
    from app.models.knowledge import KnowledgeChunk, KnowledgeDocument
    from app.services import audit_service
    file_result = await db.execute(select(FileObject).where(FileObject.id == file_id, FileObject.tenant_id == tenant_id, FileObject.status == "active"))
    file_obj = file_result.scalar_one_or_none()
    if file_obj is None:
        raise NotFoundError("File not found")
    text = extract_text(file_obj)
    chunks = chunk_text(text)
    if not chunks:
        raise ValidationAppError("File contains no extractable text")
    doc_result = await db.execute(select(KnowledgeDocument).where(KnowledgeDocument.tenant_id == tenant_id, KnowledgeDocument.file_id == file_id))
    document = doc_result.scalar_one_or_none()
    if document is None:
        document = KnowledgeDocument(tenant_id=tenant_id, file_id=file_id, status="pending")
        db.add(document)
        await db.flush()
    else:
        await db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_id == document.id))
        document.status = "pending"
        document.error_message = None
    try:
        embeddings: list[list[float]] = []
        batch_size = 32
        for start in range(0, len(chunks), batch_size):
            embeddings.extend(await embed_texts(chunks[start:start + batch_size]))
        for index, (content, embedding) in enumerate(zip(chunks, embeddings)):
            db.add(KnowledgeChunk(tenant_id=tenant_id, document_id=document.id, chunk_index=index, content=content, embedding=embedding))
        document.status = "indexed"
        document.chunk_count = len(chunks)
        document.embedding_model = "deterministic-certification" if settings.e2e_deterministic_embeddings else settings.ai_embedding_model
        await db.flush()
    except Exception as exc:
        document.status = "failed"
        document.error_message = str(exc)[:2000]
        await db.flush()
        raise
    await audit_service.record(db, action="knowledge.indexed", actor_type="user", actor_id=actor_id, tenant_id=tenant_id, resource_type="knowledge_document", resource_id=document.id, request_id=request_id_var.get(), metadata={"file_id": str(file_id), "chunk_count": len(chunks), "embedding_model": document.embedding_model})
    await db.refresh(document)
    return document


def cosine_similarity(a: list[float], b: list[float]) -> float:
    if len(a) != len(b) or not a:
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(x * x for x in b))
    return dot / (na * nb) if na and nb else 0.0


async def search(db: AsyncSession, *, tenant_id: uuid.UUID, query: str, top_k: int = 5) -> list[dict]:
    from app.models.file import FileObject
    from app.models.knowledge import KnowledgeChunk, KnowledgeDocument
    if not query.strip():
        raise ValidationAppError("Query must not be empty")
    top_k = min(max(top_k, 1), 20)
    query_embedding = (await embed_texts([query]))[0]
    result = await db.execute(select(KnowledgeChunk, KnowledgeDocument, FileObject).join(KnowledgeDocument, KnowledgeDocument.id == KnowledgeChunk.document_id).join(FileObject, FileObject.id == KnowledgeDocument.file_id).where(KnowledgeChunk.tenant_id == tenant_id, KnowledgeDocument.tenant_id == tenant_id, KnowledgeDocument.status == "indexed", FileObject.tenant_id == tenant_id, FileObject.status == "active"))
    scored = [(cosine_similarity(query_embedding, chunk.embedding), chunk, document, file_obj) for chunk, document, file_obj in result.all()]
    scored.sort(key=lambda item: item[0], reverse=True)
    return [{"chunk_id": str(chunk.id), "document_id": str(chunk.document_id), "file_id": str(file_obj.id), "filename": file_obj.filename, "chunk_index": chunk.chunk_index, "score": round(score, 6), "content": chunk.content} for score, chunk, document, file_obj in scored[:top_k]]
